"""
DOCX extractor using python-docx.

Parses Word documents for structured text, headings-based title
detection, and metadata extraction from core document properties.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.services.extraction.base import BaseExtractor, ExtractionResult

_CASE_NUMBER_PATTERN = re.compile(
    r"(?:case|file|ref(?:erence)?|docket|matter)\s*(?:#|no\.?|number)?\s*[:.]?\s*"
    r"([A-Z0-9][\w\-/.]{2,20})",
    re.IGNORECASE,
)


class DocxExtractor(BaseExtractor):
    """Extract text and metadata from .docx files via python-docx."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    async def extract(self, file_path: Path) -> ExtractionResult:
        result = ExtractionResult()

        try:
            doc = DocxDocument(str(file_path))
        except PackageNotFoundError:
            result.errors.append("Invalid or corrupted .docx file.")
            return result
        except Exception as exc:
            result.errors.append(f"Failed to open DOCX: {exc}")
            return result

        try:
            # ── Full text extraction ───────────────────────────────
            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # ── Tables (often contain contract terms, budgets) ─────
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            result.raw_text = "\n".join(paragraphs)
            result.page_count = self._estimate_page_count(result.raw_text)

            # ── Title: use first Heading style, else first paragraph
            result.title = self._detect_title(doc, paragraphs)

            # ── Case number ────────────────────────────────────────
            match = _CASE_NUMBER_PATTERN.search(result.raw_text[:5000])
            if match:
                result.case_number = match.group(1).strip()

            # ── Core properties (author, created date, etc.) ───────
            props = doc.core_properties
            if props.created:
                result.date = props.created
            if props.author:
                result.metadata["author"] = props.author
            if props.title:
                result.metadata["doc_title"] = props.title
                # Prefer the official title property over heuristic
                if not result.title:
                    result.title = props.title

        except Exception as exc:
            result.errors.append(f"Extraction error: {exc}")

        return result

    @staticmethod
    def _detect_title(
        doc: DocxDocument,
        paragraphs: list[str],
    ) -> str | None:
        """Prefer heading-styled paragraphs; fall back to first paragraph."""
        for para in doc.paragraphs:
            if (
                para.style
                and para.style.name
                and "heading" in para.style.name.lower()
            ):
                text = para.text.strip()
                if text and len(text) > 3:
                    return text[:500]
        return paragraphs[0][:500] if paragraphs else None

    @staticmethod
    def _estimate_page_count(text: str) -> int:
        """Rough estimate: ~3,000 characters per page."""
        return max(1, len(text) // 3000)
